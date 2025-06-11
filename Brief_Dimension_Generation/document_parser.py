import os
import sys
from pathlib import Path
from typing import Dict, List, Optional
import json

# Document processing imports
try:
    import pdfplumber
    import docx
    from docx import Document
    PDF_AVAILABLE = True
    DOCX_AVAILABLE = True
except ImportError as e:
    print(f"Warning: Document processing libraries not available: {e}")
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfplumber
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text as string
    """
    if not PDF_AVAILABLE:
        raise ImportError("pdfplumber not available. Install with: pip install pdfplumber")
    
    try:
        with pdfplumber.open(file_path) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            return full_text.strip()
            
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text as string
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Install with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        full_text = "\n".join(text_parts)
        return full_text.strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_document(file_path: str) -> str:
    """
    Extract text from document based on file extension
    
    Args:
        file_path: Path to document file
        
    Returns:
        Extracted text as string
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif extension in ['.docx', '.doc']:
        if extension == '.doc':
            print(f"Warning: .doc files not fully supported. Consider converting {file_path.name} to .docx")
        return extract_text_from_docx(str(file_path))
    else:
        raise ValueError(f"Unsupported file format: {extension}. Supported: .pdf, .docx")

def get_available_brief_files(brief_folder: str = "files/brief") -> Dict[str, str]:
    """
    Get all available brief files from the brief folder
    
    Args:
        brief_folder: Path to brief folder
        
    Returns:
        Dictionary mapping brief names to file paths
    """
    brief_folder = Path(brief_folder)
    
    if not brief_folder.exists():
        raise FileNotFoundError(f"Brief folder not found: {brief_folder}")
    
    brief_files = {}
    supported_extensions = ['.pdf', '.docx', '.doc']
    
    for file_path in brief_folder.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            # Create brief name from filename (without extension)
            brief_name = file_path.stem.lower().replace(' ', '_').replace('-', '_')
            brief_files[brief_name] = str(file_path)
    
    return brief_files

def load_brief_content(brief_name: str, brief_folder: str = "files/brief") -> str:
    """
    Load content from a specific brief file
    
    Args:
        brief_name: Name of the brief (filename without extension)
        brief_folder: Path to brief folder
        
    Returns:
        Brief content as plain text string
    """
    available_briefs = get_available_brief_files(brief_folder)
    
    if brief_name not in available_briefs:
        available_names = list(available_briefs.keys())
        raise ValueError(f"Brief '{brief_name}' not found. Available briefs: {available_names}")
    
    file_path = available_briefs[brief_name]
    print(f"Loading brief content from: {file_path}")
    
    content = extract_text_from_document(file_path)
    
    if not content:
        raise ValueError(f"No content extracted from {file_path}")
    
    print(f"Extracted {len(content)} characters from {brief_name}")
    return content

def load_all_brief_contents(brief_folder: str = "files/brief") -> Dict[str, str]:
    """
    Load content from all available brief files
    
    Args:
        brief_folder: Path to brief folder
        
    Returns:
        Dictionary mapping brief names to their content
    """
    available_briefs = get_available_brief_files(brief_folder)
    brief_contents = {}
    
    print(f"Found {len(available_briefs)} brief files:")
    for brief_name, file_path in available_briefs.items():
        print(f"  • {brief_name}: {Path(file_path).name}")
    
    for brief_name, file_path in available_briefs.items():
        try:
            content = load_brief_content(brief_name, brief_folder)
            brief_contents[brief_name] = content
        except Exception as e:
            print(f"Error loading {brief_name}: {e}")
            brief_contents[brief_name] = ""
    
    successful_loads = sum(1 for content in brief_contents.values() if content)
    print(f"\nSuccessfully loaded {successful_loads}/{len(available_briefs)} briefs")
    
    return brief_contents

def list_available_briefs(brief_folder: str = "files/brief") -> None:
    """
    Display all available brief files
    
    Args:
        brief_folder: Path to brief folder
    """
    try:
        available_briefs = get_available_brief_files(brief_folder)
        
        if not available_briefs:
            print(f"No brief files found in {brief_folder}")
            return
        
        print(f"Available brief files in {brief_folder}:")
        for brief_name, file_path in available_briefs.items():
            file_size = Path(file_path).stat().st_size
            file_size_mb = file_size / (1024 * 1024)
            print(f"  • {brief_name}: {Path(file_path).name} ({file_size_mb:.1f} MB)")
            
    except Exception as e:
        print(f"Error listing briefs: {e}")

# Test function
def test_document_parser():
    """
    Test the document parser with available files
    """
    print("Testing Document Parser...")
    print("=" * 50)
    
    try:
        # List available briefs
        list_available_briefs()
        
        # Load all briefs
        brief_contents = load_all_brief_contents()
        
        # Display summary
        print(f"\nBrief Loading Summary:")
        for brief_name, content in brief_contents.items():
            if content:
                print(f"  ✓ {brief_name}: {len(content)} characters")
            else:
                print(f"  ✗ {brief_name}: Failed to load")
                
    except Exception as e:
        print(f"Test failed: {e}")

if __name__ == "__main__":
    test_document_parser() 